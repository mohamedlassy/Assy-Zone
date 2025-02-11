import google.generativeai as genai
from docx import Document
from tkinter import Tk, Label, Entry, Button, messagebox, StringVar
from tkinter.ttk import Progressbar, Style, Combobox
import os
from datetime import datetime
import requests

# تكوين مفاتيح API
GENAI_API_KEY = 'AIzaSyDvN7DZ6ETx-37bYnbWMNbFKWo2TO6lGjI'
genai.configure(api_key=GENAI_API_KEY)

# مفتاح Google Custom Search API الذي زودتني به
GOOGLE_SEARCH_API_KEY = 'AIzaSyAN9T3joIAva5jWbUxAdvtgvhponcvEED4'
# تحديث معرّف محرك البحث المخصص إلى القيمة المستخرجة من الكود الذي زودتني به
GOOGLE_SEARCH_ENGINE_ID = 'a048839a17d7b404c'

# إعداد نافذة البرنامج
window = Tk()
window.title("Script AI")
window.geometry("600x550")
window.resizable(False, False)
window.configure(bg="#2E3B4E")  # خلفية زرقاء داكنة أنيقة

# إعداد الأنماط لشريط التقدم
style = Style(window)
style.theme_use("default")
style.configure("TProgressbar", thickness=20)

# عنوان البرنامج
title_label = Label(window, text="Smart Writer", font=("Arial", 24, "bold"), fg="white", bg="#2E3B4E")
title_label.pack(pady=10)

# قائمة اختيار اللغة
language_label = Label(window, text="اختر اللغة / Choose Language:", font=("Arial", 14), bg="#2E3B4E", fg="white")
language_label.pack(pady=5)
language_var = StringVar()
language_var.set("العربية")  # اللغة الافتراضية
language_combobox = Combobox(window, textvariable=language_var, values=["العربية", "English"], state="readonly")
language_combobox.pack(pady=10)

# إدخال عنوان الموضوع
topic_label = Label(window, text="أدخل عنوان الموضوع / Enter the topic title:", font=("Arial", 14), bg="#2E3B4E",
                    fg="white")
topic_label.pack(pady=5)
topic_entry = Entry(window, width=50)
topic_entry.pack(pady=10)

# المسار الافتراضي لحفظ الملفات (سطح المكتب)
save_path = os.path.join(os.path.expanduser("~"), "Desktop")
save_path_label = Label(window, text=f"المسار الحالي / Current Path: {save_path}", font=("Arial", 10), bg="#2E3B4E",
                        fg="white")
save_path_label.pack(pady=5)

# شريط التقدم
progress = Progressbar(window, length=400, mode="determinate", style="TProgressbar")
progress.pack(pady=10)


def update_progress_color(value):
    if value < 20:
        style.configure("TProgressbar", troughcolor="white", background="gray")
    elif value < 40:
        style.configure("TProgressbar", troughcolor="white", background="yellow")
    elif value < 60:
        style.configure("TProgressbar", troughcolor="white", background="orange")
    elif value < 80:
        style.configure("TProgressbar", troughcolor="white", background="red")
    elif value < 100:
        style.configure("TProgressbar", troughcolor="white", background="green")
    else:
        style.configure("TProgressbar", troughcolor="white", background="blue")


# دالة لإجراء استدعاء إلى الذكاء الاصطناعي لتوليد نص بناءً على الموجه
def generate_text(prompt):
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"خطأ أثناء التوليد / Error during generation: {e}")
        return None


# دالة للبحث في جوجل باستخدام Google Custom Search API
def google_search(query, num_results=3):
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": GOOGLE_SEARCH_API_KEY,
        "cx": GOOGLE_SEARCH_ENGINE_ID,
        "q": query,
        "num": num_results
    }
    response = requests.get(url, params=params)
    results = []
    if response.status_code == 200:
        data = response.json()
        items = data.get("items", [])
        for item in items:
            results.append({
                "title": item.get("title"),
                "link": item.get("link"),
                "snippet": item.get("snippet")
            })
    else:
        messagebox.showerror("خطأ / Error", f"فشل البحث في جوجل: {response.status_code}")
    return results


# دالة لحفظ المستند بصيغة DOCX على سطح المكتب
def save_document(doc, title):
    try:
        filename = os.path.join(save_path, f"{title}_{datetime.now().strftime('%Y-%m-%d')}.docx")
        doc.save(filename)
        return filename
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"فشل حفظ الملف / Failed to save file: {e}")
        return None


# دالة لتوليد مخطط/عناصر الموضوع باستخدام الذكاء الاصطناعي
def generate_outline(topic, language):
    if language == "العربية":
        prompt = f"اعطني قائمة من العناصر الرئيسية التي يجب تناولها في موضوع '{topic}'. اكتب كل عنصر في سطر منفصل دون أرقام أو علامات."
    else:
        prompt = f"Provide a list of main elements to cover in an article about '{topic}', each on a new line without numbers or bullet points."
    return generate_text(prompt)


# دالة لتوليد محتوى تفصيلي لكل عنصر، مع تضمين نتائج بحث جوجل كمراجع
def generate_detail(element, language):
    # الحصول على مراجع من جوجل لكل عنصر
    references = google_search(element, num_results=3)
    # إعداد نص المراجع
    if references:
        references_text = "\n".join(
            [f"Title: {ref['title']}\nLink: {ref['link']}\nSnippet: {ref['snippet']}" for ref in references]
        )
    else:
        references_text = "لا توجد مراجع متاحة."

    if language == "العربية":
        prompt = (
            f"اكتب مقالة تفصيلية عن '{element}'. استخدم المعلومات الواردة أدناه (نتائج بحث جوجل) كمراجع واستعن بها "
            f"لتحسين جودة المحتوى:\n{references_text}\n"
            "حاول استخدام كل المراجع والطرق المتاحة للحصول على أفضل محتوى ممكن.")
    else:
        prompt = (
            f"Write a detailed article section about '{element}'. Use the following Google search results as references to enhance the content:\n"
            f"{references_text}\n"
            "Employ all available references and methods to produce the best possible content.")
    return generate_text(prompt)


# العملية الرئيسية: تقسيم الموضوع إلى عناصر، توليد محتوى لكل عنصر باستخدام مراجع بحث جوجل، ومن ثم حفظ الملف
def generate_and_save():
    topic = topic_entry.get().strip()
    language = language_var.get()

    if not topic:
        messagebox.showerror("خطأ / Error", "يرجى إدخال عنوان الموضوع / Please enter the topic title.")
        return

    # المرحلة الأولى: توليد عناصر الموضوع
    progress["value"] = 10
    update_progress_color(10)
    status_label.config(text="جارٍ توليد عناصر الموضوع... / Generating outline...")
    window.update_idletasks()

    outline_text = generate_outline(topic, language)
    if not outline_text:
        status_label.config(text="فشل توليد العناصر / Outline generation failed.")
        return

    # تقسيم النص الناتج إلى قائمة من العناصر (يفترض كل عنصر بسطر منفصل)
    outline_items = [line.strip() for line in outline_text.splitlines() if line.strip()]
    if not outline_items:
        outline_items = [outline_text.strip()]

    progress["value"] = 30
    update_progress_color(30)
    window.update_idletasks()

    # المرحلة الثانية: توليد المحتوى التفصيلي لكل عنصر باستخدام مراجع بحث جوجل
    details = []
    num_items = len(outline_items)
    for i, element in enumerate(outline_items):
        status_label.config(text=f"جارٍ توليد المحتوى للعنصر: {element}...")
        window.update_idletasks()
        detail = generate_detail(element, language)
        if not detail:
            detail = "لم يتم التوليد بنجاح."
        details.append((element, detail))
        progress["value"] = 30 + int((i + 1) * 60 / num_items)  # تحديث شريط التقدم من 30 إلى 90
        update_progress_color(progress["value"])
        window.update_idletasks()

    # المرحلة الثالثة: إنشاء وتنسيق المستند
    doc = Document()
    doc.add_heading(topic, level=1)
    for element, detail in details:
        doc.add_heading(element, level=2)
        doc.add_paragraph(detail)

    progress["value"] = 95
    update_progress_color(95)
    window.update_idletasks()

    filename = save_document(doc, topic)
    if filename:
        progress["value"] = 100
        update_progress_color(100)
        status_label.config(text=f"تم الحفظ بنجاح في / Successfully saved in: {filename}")
        messagebox.showinfo("نجاح / Success",
                            f"تم إنشاء الملف وحفظه بصيغة DOCX / File saved successfully as DOCX:\n{filename}")
    else:
        status_label.config(text="فشل الحفظ / Save failed.")


# زر بدء العملية
generate_button = Button(window, text="توليد الموضوع وحفظه / Generate and Save", command=generate_and_save,
                         font=("Arial", 14), bg="#4CAF50", fg="white")
generate_button.pack(pady=20)

# حالة العملية
status_label = Label(window, text="الحالة / Status: جاهز / Ready", font=("Arial", 12), fg="white", bg="#2E3B4E")
status_label.pack(pady=10)

window.mainloop()

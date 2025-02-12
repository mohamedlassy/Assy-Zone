import os
import logging
from datetime import datetime
from io import BytesIO
import tkinter as tk
from tkinter import ttk, messagebox
import requests

# مكتبات التعامل مع ملفات الوورد والصيغ الأخرى
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

# مكتبات تحويل النص إلى صوت
from gtts import gTTS

# مكتبة الذكاء الاصطناعي التوليدي من جوجل
import google.generativeai as genai

# ===================== إعدادات API والمفاتيح =====================
GENAI_API_KEY = 'AIzaSyDvN7DZ6ETx-37bYnbWMNbFKWo2TO6lGjI'
GOOGLE_SEARCH_API_KEY = 'AIzaSyAN9T3joIAva5jWbUxAdvtgvhponcvEED4'
GOOGLE_SEARCH_ENGINE_ID = 'a048839a17d7b404c'  # معرّف محرك البحث المخصص
PEXELS_API_KEY = 'Afw2O4LTmWoNOps7spA9YhYAFbFVol8KfWjfSRxU059Wd8UqUXHRSAAJ'

# تكوين مكتبة Generative AI
genai.configure(api_key=GENAI_API_KEY)

# ===================== إعداد سجل الأخطاء (Logging) =====================
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)


# ===================== دوال الصور =====================
def get_pexels_images(query, num_images=2):
    """
    تبحث هذه الدالة عن صور في Pexels وترجع قائمة من الصور بصيغة BytesIO.
    """
    images = []
    url = "https://api.pexels.com/v1/search"
    headers = {"Authorization": PEXELS_API_KEY}
    params = {"query": query, "per_page": num_images}
    try:
        response = requests.get(url, headers=headers, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        photos = data.get("photos", [])
        for photo in photos:
            image_url = photo.get("src", {}).get("medium")
            if image_url:
                img_response = requests.get(image_url, timeout=10)
                img_response.raise_for_status()
                images.append(BytesIO(img_response.content))
    except Exception as e:
        logging.error(f"Pexels API Error: {e}")
        messagebox.showerror("Pexels API Error", f"Error fetching images from Pexels: {e}")
    return images


def insert_images(doc, query):
    """
    تقوم هذه الدالة بإدراج صور من Pexels داخل ملف الوورد.
    """
    images = get_pexels_images(query, num_images=2)
    if images:
        doc.add_heading("Images / الصور", level=2)
        for img in images:
            try:
                doc.add_picture(img, width=Inches(5))
            except Exception as e:
                logging.warning(f"Image Insert Warning: {e}")
                messagebox.showwarning("Image Insert Warning", f"Could not insert an image: {e}")


# ===================== دوال تحويل النص إلى صوت =====================
def generate_audio(text, title, language):
    """
    تقوم هذه الدالة بتحويل النص إلى ملف صوتي باستخدام gTTS.
    """
    tts_lang = 'ar' if language == "العربية" else 'en'
    try:
        tts = gTTS(text=text, lang=tts_lang)
        clean_title = title.replace("##", "").strip()
        audio_filename = os.path.join(save_path, f"{clean_title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.mp3")
        tts.save(audio_filename)
        logging.info(f"Audio generated: {audio_filename}")
        return audio_filename
    except Exception as e:
        logging.error(f"Audio Generation Error: {e}")
        messagebox.showerror("Audio Generation Error", f"Error generating audio with gTTS: {e}")
        return None


def generate_audio_elevenlabs(text, title, api_key, voice_id):
    """
    تقوم هذه الدالة بتحويل النص إلى صوت باستخدام واجهة ElevenLabs.
    """
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
    headers = {
        "xi-api-key": api_key,
        "Content-Type": "application/json"
    }
    payload = {
        "text": text,
        "model_id": "eleven_multilingual_v2",  # نموذج متعدد اللغات يدعم العربية
        "voice_settings": {
            "stability": 0.5,
            "similarity_boost": 0.5
        }
    }
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=15)
        if response.status_code == 200:
            clean_title = title.replace("##", "").strip()
            audio_filename = os.path.join(
                save_path,
                f"{clean_title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}_elevenlabs.mp3"
            )
            with open(audio_filename, "wb") as f:
                f.write(response.content)
            logging.info(f"ElevenLabs audio generated: {audio_filename}")
            return audio_filename
        else:
            error_message = f"Error generating audio from ElevenLabs: {response.status_code}\n{response.text}"
            logging.error(error_message)
            messagebox.showerror("ElevenLabs Error", error_message)
            return None
    except Exception as e:
        logging.error(f"ElevenLabs Exception: {e}")
        messagebox.showerror("ElevenLabs Error", f"Exception during ElevenLabs audio generation: {e}")
        return None


def convert_to_pdf(docx_filename):
    """
    تقوم هذه الدالة بتحويل ملف الوورد إلى PDF.
    """
    try:
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        convert(docx_filename, pdf_filename)
        logging.info(f"Converted to PDF: {pdf_filename}")
        return pdf_filename
    except Exception as e:
        logging.error(f"PDF Conversion Error: {e}")
        messagebox.showerror("PDF Conversion Error", f"Error converting to PDF: {e}")
        return None


# ===================== دوال البحث وتوليد المحتوى =====================
def generate_text(prompt):
    """
    ترسل هذه الدالة prompt إلى نموذج الذكاء الاصطناعي وتعيد النص المولد.
    """
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        generated = response.text.strip()
        logging.info("Article generated successfully.")
        return generated
    except Exception as e:
        logging.error(f"Text Generation Error: {e}")
        messagebox.showerror("خطأ / Error", f"Error during generation: {e}")
        return None


def google_search(query, num_results=3, language="English"):
    """
    تقوم هذه الدالة بالبحث في جوجل باستخدام API الخاص بالبحث المخصص،
    وتعيد قائمة النتائج مع دعم اللغة، وإعدادات headers وtimeout.
    """
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": GOOGLE_SEARCH_API_KEY,
        "cx": GOOGLE_SEARCH_ENGINE_ID,
        "q": query,
        "num": num_results,
        "safe": "active"
    }
    params["lr"] = "lang_ar" if language == "العربية" else "lang_en"
    headers = {"User-Agent": "SmartWriter/1.0"}

    try:
        response = requests.get(url, params=params, headers=headers, timeout=10)
        response.raise_for_status()
        data = response.json()
        items = data.get("items", [])
        results = []
        for item in items:
            results.append({
                "title": item.get("title", "No title"),
                "link": item.get("link", ""),
                "snippet": item.get("snippet", "")
            })
        logging.info(f"Google search successful for query: {query}")
        return results
    except Exception as e:
        logging.error(f"Google Search Error: {e}")
        messagebox.showerror("خطأ / Error", f"حدث خطأ أثناء بحث جوجل: {e}")
        return []


def save_document(doc, title):
    """
    تقوم هذه الدالة بحفظ ملف الوورد على سطح المكتب مع تضمين تاريخ ووقت الحفظ.
    """
    try:
        filename = os.path.join(save_path, f"{title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx")
        doc.save(filename)
        logging.info(f"DOCX saved: {filename}")
        return filename
    except Exception as e:
        logging.error(f"Save Document Error: {e}")
        messagebox.showerror("خطأ / Error", f"Failed to save file: {e}")
        return None


def generate_article(topic, language):
    """
    توليد مقال احترافي باستخدام الذكاء الاصطناعي مع تضمين مرجع بحث جوجل.
    يتم بناء prompt مفصل بناءً على المراجع المستخرجة واللغة المحددة.
    """
    references = google_search(topic, num_results=3, language=language)
    if references:
        if language == "العربية":
            references_text = "\n\n".join(
                [f"العنوان: {ref['title']}\nالرابط: {ref['link']}\nالملخص: {ref['snippet']}" for ref in references]
            )
        else:
            references_text = "\n\n".join(
                [f"Title: {ref['title']}\nLink: {ref['link']}\nSnippet: {ref['snippet']}" for ref in references]
            )
    else:
        references_text = "لا توجد مراجع متاحة." if language == "العربية" else "No references available."

    if language == "العربية":
        prompt = (
            f"اكتب مقالاً احترافيًا وشاملاً حول موضوع '{topic}'. يجب أن يتضمن المقال تحليلاً عميقاً، "
            "ورؤى مبتكرة، وأمثلة وتفاصيل دقيقة. استخدم المعلومات التالية المستخرجة من بحث جوجل كمرجع "
            "لزيادة مصداقية المحتوى:\n\n"
            f"{references_text}\n\n"
            "رتب الأفكار والفقرات بشكل منطقي لضمان تدفق سلس للمحتوى."
        )
    else:
        prompt = (
            f"Compose a professional and comprehensive article about '{topic}'. The article should include in-depth analysis, "
            "innovative insights, and detailed examples. Use the following Google search results as references to enhance the content's credibility:\n\n"
            f"{references_text}\n\n"
            "Ensure that the ideas and paragraphs are well-organized for a smooth flow of information."
        )
    return generate_text(prompt)


# ===================== إعداد واجهة المستخدم (Tkinter UI) =====================
window = tk.Tk()
window.title("Smart Writer - برنامج الكتابة الذكية")
window.geometry("750x980")
window.resizable(False, False)
window.configure(bg="#2E3B4E")

# إعداد قائمة القوائم (Menu Bar)
menu_bar = tk.Menu(window)
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Exit", command=window.quit)
menu_bar.add_cascade(label="File", menu=file_menu)
help_menu = tk.Menu(menu_bar, tearoff=0)
help_menu.add_command(label="About",
                      command=lambda: messagebox.showinfo("About", "Smart Writer\nDeveloped by Professional Developer"))
menu_bar.add_cascade(label="Help", menu=help_menu)
window.config(menu=menu_bar)

# إعداد أسلوب شريط التقدم
style = ttk.Style(window)
style.theme_use("clam")
style.configure("TProgressbar", thickness=20, background="#4CAF50", troughcolor="#D3D3D3")

# رأس التطبيق
header_frame = tk.Frame(window, bg="#2E3B4E")
header_frame.pack(pady=20)
title_label = tk.Label(header_frame, text="Smart Writer", font=("Helvetica", 36, "bold"), fg="#FFD700", bg="#2E3B4E")
title_label.pack()
subtitle_label = tk.Label(header_frame, text="برنامج الكتابة الذكية", font=("Helvetica", 18), fg="white", bg="#2E3B4E")
subtitle_label.pack()

# إطار الإدخالات العامة
input_frame = tk.Frame(window, bg="#2E3B4E")
input_frame.pack(pady=10, padx=20, fill="x")

# اختيار اللغة
lang_frame = tk.Frame(input_frame, bg="#2E3B4E")
lang_frame.pack(anchor="w", pady=10)
lang_label = tk.Label(lang_frame, text="اختر اللغة / Choose Language:", font=("Arial", 14), bg="#2E3B4E", fg="white")
lang_label.pack(side="left", padx=(0, 10))
language_var = tk.StringVar(value="العربية")
lang_combobox = ttk.Combobox(lang_frame, textvariable=language_var, values=["العربية", "English"], state="readonly",
                             font=("Arial", 12), width=12)
lang_combobox.pack(side="left")

# إدخال عنوان الموضوع
topic_frame = tk.Frame(input_frame, bg="#2E3B4E")
topic_frame.pack(anchor="w", pady=10)
topic_label = tk.Label(topic_frame, text="أدخل عنوان الموضوع / Enter Topic Title:", font=("Arial", 14), bg="#2E3B4E",
                       fg="white")
topic_label.pack(side="left", padx=(0, 10))
topic_entry = tk.Entry(topic_frame, width=40, font=("Arial", 12))
topic_entry.pack(side="left")


def show_topic_context_menu(event):
    try:
        topic_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        topic_context_menu.grab_release()


topic_context_menu = tk.Menu(window, tearoff=0)
topic_context_menu.add_command(label="Paste", command=lambda: topic_entry.insert(tk.INSERT, window.clipboard_get()))
topic_entry.bind("<Button-3>", show_topic_context_menu)

# اختيار نوع المخرجات
output_frame = tk.Frame(input_frame, bg="#2E3B4E")
output_frame.pack(anchor="w", pady=10)
output_label = tk.Label(output_frame, text="اختر نوع المخرجات / Select Output Type:", font=("Arial", 14), bg="#2E3B4E",
                        fg="white")
output_label.pack(side="left", padx=(0, 10))
output_options = ["حفظ كملف وورد (DOCX)", "تحويل إلى PDF", "تحويل إلى صوت (MP3)", "تحويل إلى صوت (ElevenLabs)"]
output_var = tk.StringVar(value=output_options[0])
output_combobox = ttk.Combobox(output_frame, textvariable=output_var, values=output_options, state="readonly",
                               font=("Arial", 12), width=30)
output_combobox.pack(side="left")

# بيانات ElevenLabs (تظهر فقط عند الحاجة)
elevenlabs_frame = tk.Frame(input_frame, bg="#2E3B4E")
elevenlabs_api_label = tk.Label(elevenlabs_frame, text="ElevenLabs API Key:", font=("Arial", 12), bg="#2E3B4E",
                                fg="white")
elevenlabs_api_label.pack(side="left", padx=(0, 5))
elevenlabs_api_entry = tk.Entry(elevenlabs_frame, width=30, font=("Arial", 12))
elevenlabs_api_entry.pack(side="left", padx=(0, 10))


def show_api_context_menu(event):
    try:
        api_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        api_context_menu.grab_release()


api_context_menu = tk.Menu(window, tearoff=0)
api_context_menu.add_command(label="Paste",
                             command=lambda: elevenlabs_api_entry.insert(tk.INSERT, window.clipboard_get()))
elevenlabs_api_entry.bind("<Button-3>", show_api_context_menu)

# خيارات ملف الوورد الوسيط (للبعض من المخرجات)
doc_option_frame = tk.Frame(input_frame, bg="#2E3B4E")
doc_option_label = tk.Label(doc_option_frame, text="ملف الوورد الوسيط:", font=("Arial", 14), bg="#2E3B4E", fg="white")
doc_option_label.pack(side="left", padx=(0, 10))
doc_option_var = tk.StringVar(value="حذف ملف الوورد")
doc_option_combobox = ttk.Combobox(doc_option_frame, textvariable=doc_option_var,
                                   values=["حذف ملف الوورد", "الابقاء على ملف الوورد"], state="readonly",
                                   font=("Arial", 12), width=30)
doc_option_combobox.pack(side="left")


def update_output_options(*args):
    if output_var.get() == "تحويل إلى صوت (ElevenLabs)":
        elevenlabs_frame.pack(anchor="w", pady=5)
    else:
        elevenlabs_frame.pack_forget()
    if output_var.get() in ["تحويل إلى PDF", "تحويل إلى صوت (MP3)", "تحويل إلى صوت (ElevenLabs)"]:
        doc_option_frame.pack(anchor="w", pady=5)
    else:
        doc_option_frame.pack_forget()


output_var.trace_add("write", update_output_options)

# تحديد مسار الحفظ (سطح المكتب)
save_path = os.path.join(os.path.expanduser("~"), "Desktop")
save_path_info = tk.Label(window, text=f"سيتم الحفظ في: {save_path}", font=("Arial", 10), bg="#2E3B4E", fg="white")
save_path_info.pack(pady=5)

# شريط التقدم والحالة
progress = ttk.Progressbar(window, length=600, mode="determinate", style="TProgressbar")
progress.pack(pady=15)
status_label = tk.Label(window, text="الحالة / Status: جاهز / Ready", font=("Arial", 12), bg="#2E3B4E", fg="white")
status_label.pack(pady=5)


def update_progress_color(value):
    """
    تقوم هذه الدالة بتحديث لون شريط التقدم بناءً على القيمة الحالية.
    """
    if value < 20:
        style.configure("TProgressbar", background="gray")
    elif value < 40:
        style.configure("TProgressbar", background="yellow")
    elif value < 60:
        style.configure("TProgressbar", background="orange")
    elif value < 80:
        style.configure("TProgressbar", background="red")
    elif value < 100:
        style.configure("TProgressbar", background="green")
    else:
        style.configure("TProgressbar", background="blue")


# ===================== العملية الرئيسية =====================
def generate_and_save():
    topic = topic_entry.get().strip()
    language = language_var.get()
    selected_output = output_var.get()

    if not topic:
        messagebox.showerror("خطأ / Error", "يرجى إدخال عنوان الموضوع / Please enter the topic title.")
        return

    progress["value"] = 10
    update_progress_color(10)
    status_label.config(text="جارٍ توليد المقالة... / Generating article...")
    window.update_idletasks()

    article_text = generate_article(topic, language)
    if not article_text:
        status_label.config(text="فشل توليد المقالة / Article generation failed.")
        return

    # إزالة العلامات الزائدة من بداية السطر الأول إذا وُجدت
    lines = article_text.splitlines()
    if lines and lines[0].startswith("##"):
        lines[0] = lines[0].lstrip("# ").strip()
        article_text = "\n".join(lines)

    progress["value"] = 50
    update_progress_color(50)
    window.update_idletasks()

    # إنشاء وثيقة الوورد وتنسيق النص بناءً على اللغة
    doc = Document()
    if language == "العربية":
        main_title_font = "Amiri"
        body_font = "Amiri"
        main_title_alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        main_title_font = "Calibri"
        body_font = "Calibri"
        main_title_alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_alignment = WD_ALIGN_PARAGRAPH.LEFT

    # إضافة العنوان الرئيسي
    main_title = doc.add_heading(topic, level=1)
    main_title.alignment = main_title_alignment
    for run in main_title.runs:
        run.font.name = main_title_font
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 255)

    # تقسيم النص إلى فقرات والعناوين الفرعية (الأسطر التي تبدأ بـ "##")
    for line in article_text.splitlines():
        if line.strip().startswith("##"):
            subheading_text = line.strip().lstrip("#").strip()
            subheading = doc.add_heading(subheading_text, level=2)
            subheading.alignment = body_alignment
            for run in subheading.runs:
                run.font.name = body_font
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            para = doc.add_paragraph(line)
            para.alignment = body_alignment
            for run in para.runs:
                run.font.name = body_font
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

    # إدراج الصور المستخرجة من Pexels
    insert_images(doc, topic)

    # حفظ ملف الوورد الوسيط
    docx_filename = save_document(doc, topic)
    if not docx_filename:
        status_label.config(text="فشل حفظ ملف الوورد / Failed to save DOCX.")
        return

    # التعامل مع نوع المخرجات المختار
    if selected_output == "حفظ كملف وورد (DOCX)":
        output_file = docx_filename
    elif selected_output == "تحويل إلى PDF":
        output_file = convert_to_pdf(docx_filename)
        if output_file and doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                logging.warning(f"Unable to remove DOCX file: {e}")
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    elif selected_output == "تحويل إلى صوت (MP3)":
        output_file = generate_audio(article_text, topic, language)
        if doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                logging.warning(f"Unable to remove DOCX file: {e}")
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    elif selected_output == "تحويل إلى صوت (ElevenLabs)":
        api_key = elevenlabs_api_entry.get().strip()
        if not api_key:
            messagebox.showerror("خطأ / Error", "يرجى إدخال ElevenLabs API Key.")
            return
        voice_id = "XrExE9yKIg1WjnnlVkGX"  # Voice ID الافتراضي
        output_file = generate_audio_elevenlabs(article_text, topic, api_key, voice_id)
        if doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                logging.warning(f"Unable to remove DOCX file: {e}")
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    else:
        output_file = None

    progress["value"] = 100
    update_progress_color(100)
    window.update_idletasks()

    if output_file:
        status_label.config(text=f"تم الحفظ بنجاح / Successfully saved: {output_file}")
        messagebox.showinfo("نجاح / Success", f"تم إنشاء الملف:\n{output_file}")
    else:
        status_label.config(text="لم يتم إنشاء أي ملف.")


# زر بدء العملية
generate_button = tk.Button(window, text="توليد المقالة وحفظها / Generate and Save", command=generate_and_save,
                            font=("Arial", 14, "bold"), bg="#4CAF50", fg="white", padx=10, pady=5)
generate_button.pack(pady=20)

# ===================== بدء حلقة التطبيق الرئيسية =====================
if __name__ == '__main__':
    window.mainloop()

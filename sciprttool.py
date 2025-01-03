import google.generativeai as genai
from docx import Document
from datetime import datetime

API_KEY = 'ضع مفتاحك هنا'

# Configure the API
genai.configure(api_key=API_KEY)

# Define the prompt
prompt = (
    "اكتب لي سكربت لفيديو قصير مدته 60 ثانية لقناة يوتيوب متخصصة في الحقائق الرياضية. يجب أن يكون السكربت جذابًا ومثيرًا، يركز على قصة أو حقيقة عن لاعب كرة قدم، ويمزج بين الإثارة والواقعية. استخدم مثالًا مشهورًا مثل زين الدين زيدان، وذكر كيف غير مسيرة مباراة مهمة بطريقة غير متوقعة. اجعل السكربت يتضمن دعوة للمشاهدين للإعجاب بالفيديو، والتعليق، والاشتراك في القناة لمزيد من الفيديوهات. اجعل الأسلوب حيويًا وشيقًا، مع التركيز على جذب الانتباه بسرعة. "

)

# Initialize the chat model
model = genai.GenerativeModel('gemini-1.5-flash')
chat = model.start_chat(history=[])

# Send the prompt to the model and get the response
response = chat.send_message(prompt)

# Create a new Word document
doc = Document()

# Add a title and the response to the document
doc.add_heading('Script for YouTube Video', 0)
doc.add_paragraph(response.text)

# Generate a filename based on the current date
filename = datetime.now().strftime("script.docx")

# Save the document
doc.save(filename)

print(f"Script saved as {filename}")

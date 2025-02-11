import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import ttk, messagebox
import os
from datetime import datetime
import requests
from io import BytesIO

# مكتبات الميزات الجديدة:
from gtts import gTTS  # لتحويل النص إلى كلام باستخدام gTTS
from docx2pdf import convert  # لتحويل DOCX إلى PDF

# --------------------------- إعدادات وتكوينات API ---------------------------
GENAI_API_KEY = 'AIzaSyDvN7DZ6ETx-37bYnbWMNbFKWo2TO6lGjI'
genai.configure(api_key=GENAI_API_KEY)

GOOGLE_SEARCH_API_KEY = 'AIzaSyAN9T3joIAva5jWbUxAdvtgvhponcvEED4'
GOOGLE_SEARCH_ENGINE_ID = 'a048839a17d7b404c'  # معرّف محرك البحث المخصص

PEXELS_API_KEY = 'Afw2O4LTmWoNOps7spA9YhYAFbFVol8KfWjfSRxU059Wd8UqUXHRSAAJ'


# --------------------------- دوال البحث عن الصور ---------------------------
def get_pexels_images(query, num_images=2):
    images = []
    url = "https://api.pexels.com/v1/search"
    headers = {"Authorization": PEXELS_API_KEY}
    params = {"query": query, "per_page": num_images}
    try:
        response = requests.get(url, headers=headers, params=params)
        if response.status_code == 200:
            data = response.json()
            photos = data.get("photos", [])
            for photo in photos:
                image_url = photo.get("src", {}).get("medium")
                if image_url:
                    img_response = requests.get(image_url)
                    if img_response.status_code == 200:
                        images.append(BytesIO(img_response.content))
        else:
            messagebox.showwarning("Pexels API Warning", f"Pexels API returned status code {response.status_code}")
    except Exception as e:
        messagebox.showerror("Pexels API Error", f"Error fetching images from Pexels: {e}")
    return images


def insert_images(doc, query):
    images = get_pexels_images(query, num_images=2)
    if images:
        doc.add_heading("Images / الصور", level=2)
        for img in images:
            try:
                doc.add_picture(img, width=Inches(5))
            except Exception as e:
                messagebox.showwarning("Image Insert Warning", f"Could not insert an image: {e}")


# --------------------------- دوال الميزات الجديدة ---------------------------
def generate_audio(text, title, language):
    tts_lang = 'ar' if language == "العربية" else 'en'
    try:
        tts = gTTS(text=text, lang=tts_lang)
        clean_title = title.replace("##", "")
        audio_filename = os.path.join(save_path, f"{clean_title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.mp3")
        tts.save(audio_filename)
        return audio_filename
    except Exception as e:
        messagebox.showerror("Audio Generation Error", f"Error generating audio with gTTS: {e}")
        return None


def generate_audio_elevenlabs(text, title, api_key, voice_id):
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
    headers = {"xi-api-key": api_key, "Content-Type": "application/json"}
    payload = {"text": text, "voice_settings": {"stability": 0.5, "similarity_boost": 0.5}}
    try:
        response = requests.post(url, json=payload, headers=headers)
        if response.status_code == 200:
            clean_title = title.replace("##", "")
            audio_filename = os.path.join(save_path,
                                          f"{clean_title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}_elevenlabs.mp3")
            with open(audio_filename, "wb") as f:
                f.write(response.content)
            return audio_filename
        else:
            messagebox.showerror("ElevenLabs Error",
                                 f"Error generating audio from ElevenLabs: {response.status_code}\n{response.text}")
            return None
    except Exception as e:
        messagebox.showerror("ElevenLabs Error", f"Exception during ElevenLabs audio generation: {e}")
        return None


def convert_to_pdf(docx_filename):
    try:
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        convert(docx_filename, pdf_filename)
        return pdf_filename
    except Exception as e:
        messagebox.showerror("PDF Conversion Error", f"Error converting to PDF: {e}")
        return None


# --------------------------- إعداد واجهة المستخدم ---------------------------
window = tk.Tk()
window.title("Smart Writer - برنامج الكتابة الذكية")
window.geometry("750x980")
window.resizable(False, False)
window.configure(bg="#2E3B4E")

menu_bar = tk.Menu(window)
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Exit", command=window.quit)
menu_bar.add_cascade(label="File", menu=file_menu)
help_menu = tk.Menu(menu_bar, tearoff=0)
help_menu.add_command(label="About",
                      command=lambda: messagebox.showinfo("About", "Smart Writer\nDeveloped by Professional Developer"))
menu_bar.add_cascade(label="Help", menu=help_menu)
window.config(menu=menu_bar)

style = ttk.Style(window)
style.theme_use("clam")
style.configure("TProgressbar", thickness=20, background="#4CAF50", troughcolor="#D3D3D3")

# الرأس (Header)
header_frame = tk.Frame(window, bg="#2E3B4E")
header_frame.pack(pady=20)
title_label = tk.Label(header_frame, text="Smart Writer", font=("Helvetica", 36, "bold"), fg="#FFD700", bg="#2E3B4E")
title_label.pack()
subtitle_label = tk.Label(header_frame, text="برنامج الكتابة الذكية", font=("Helvetica", 18), fg="white", bg="#2E3B4E")
subtitle_label.pack()

# الإطار العام للإدخالات (لا توجد قوائم لاختيار التنسيق الآن)
input_frame = tk.Frame(window, bg="#2E3B4E")
input_frame.pack(pady=10, padx=20, fill="x")

# اختيار اللغة
lang_frame = tk.Frame(input_frame, bg="#2E3B4E")
lang_frame.pack(anchor="w", pady=10)
lang_label = tk.Label(lang_frame, text="اختر اللغة / Choose Language:", font=("Arial", 14), bg="#2E3B4E", fg="white")
lang_label.pack(side="left", padx=(0, 10))
language_var = tk.StringVar(value="العربية")
lang_combobox = ttk.Combobox(lang_frame, textvariable=language_var, values=["العربية", "English"], state="readonly",
                             font=("Arial", 12), width=12)
lang_combobox.pack(side="left")

# عنوان الموضوع
topic_frame = tk.Frame(input_frame, bg="#2E3B4E")
topic_frame.pack(anchor="w", pady=10)
topic_label = tk.Label(topic_frame, text="أدخل عنوان الموضوع / Enter Topic Title:", font=("Arial", 14), bg="#2E3B4E",
                       fg="white")
topic_label.pack(side="left", padx=(0, 10))
topic_entry = tk.Entry(topic_frame, width=40, font=("Arial", 12))
topic_entry.pack(side="left")


def show_topic_context_menu(event):
    try:
        topic_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        topic_context_menu.grab_release()


topic_context_menu = tk.Menu(window, tearoff=0)
topic_context_menu.add_command(label="Paste", command=lambda: topic_entry.insert(tk.INSERT, window.clipboard_get()))
topic_entry.bind("<Button-3>", show_topic_context_menu)

# لا توجد قوائم اختيار للتنسيق؛ يتم التعيين الثابت

# اختيار نوع المخرجات (تبقى هذه القائمة كما هي)
output_frame = tk.Frame(input_frame, bg="#2E3B4E")
output_frame.pack(anchor="w", pady=10)
output_label = tk.Label(output_frame, text="اختر نوع المخرجات / Select Output Type:", font=("Arial", 14), bg="#2E3B4E",
                        fg="white")
output_label.pack(side="left", padx=(0, 10))
output_options = ["حفظ كملف وورد (DOCX)", "تحويل إلى PDF", "تحويل إلى صوت (MP3)", "تحويل إلى صوت (ElevenLabs)"]
output_var = tk.StringVar(value=output_options[0])
output_combobox = ttk.Combobox(output_frame, textvariable=output_var, values=output_options, state="readonly",
                               font=("Arial", 12), width=30)
output_combobox.pack(side="left")

# بيانات ElevenLabs (تبقى كما هي)
elevenlabs_frame = tk.Frame(input_frame, bg="#2E3B4E")
elevenlabs_api_label = tk.Label(elevenlabs_frame, text="ElevenLabs API Key:", font=("Arial", 12), bg="#2E3B4E",
                                fg="white")
elevenlabs_api_label.pack(side="left", padx=(0, 5))
elevenlabs_api_entry = tk.Entry(elevenlabs_frame, width=30, font=("Arial", 12))
elevenlabs_api_entry.pack(side="left", padx=(0, 10))


def show_api_context_menu(event):
    try:
        api_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        api_context_menu.grab_release()


api_context_menu = tk.Menu(window, tearoff=0)
api_context_menu.add_command(label="Paste",
                             command=lambda: elevenlabs_api_entry.insert(tk.INSERT, window.clipboard_get()))
elevenlabs_api_entry.bind("<Button-3>", show_api_context_menu)
elevenlabs_voice_label = tk.Label(elevenlabs_frame, text="Voice ID:", font=("Arial", 12), bg="#2E3B4E", fg="white")
elevenlabs_voice_label.pack(side="left", padx=(0, 5))
elevenlabs_voice_entry = tk.Entry(elevenlabs_frame, width=20, font=("Arial", 12))
elevenlabs_voice_entry.pack(side="left")


def show_voice_context_menu(event):
    try:
        voice_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        voice_context_menu.grab_release()


voice_context_menu = tk.Menu(window, tearoff=0)
voice_context_menu.add_command(label="Paste",
                               command=lambda: elevenlabs_voice_entry.insert(tk.INSERT, window.clipboard_get()))
elevenlabs_voice_entry.bind("<Button-3>", show_voice_context_menu)

# خيارات ملف الوورد الوسيط (تبقى كما هي)
doc_option_frame = tk.Frame(input_frame, bg="#2E3B4E")
doc_option_label = tk.Label(doc_option_frame, text="ملف الوورد الوسيط:", font=("Arial", 14), bg="#2E3B4E", fg="white")
doc_option_label.pack(side="left", padx=(0, 10))
doc_option_var = tk.StringVar(value="حذف ملف الوورد")
doc_option_combobox = ttk.Combobox(doc_option_frame, textvariable=doc_option_var,
                                   values=["حذف ملف الوورد", "الابقاء على ملف الوورد"], state="readonly",
                                   font=("Arial", 12), width=30)
doc_option_combobox.pack(side="left")


def update_output_options(*args):
    if output_var.get() == "تحويل إلى صوت (ElevenLabs)":
        elevenlabs_frame.pack(anchor="w", pady=5)
    else:
        elevenlabs_frame.pack_forget()
    if output_var.get() in ["تحويل إلى PDF", "تحويل إلى صوت (MP3)", "تحويل إلى صوت (ElevenLabs)"]:
        doc_option_frame.pack(anchor="w", pady=5)
    else:
        doc_option_frame.pack_forget()


output_var.trace_add("write", update_output_options)

# مسار الحفظ (سطح المكتب)
save_path = os.path.join(os.path.expanduser("~"), "Desktop")
save_path_info = tk.Label(window, text=f"سيتم الحفظ في: {save_path}", font=("Arial", 10), bg="#2E3B4E", fg="white")
save_path_info.pack(pady=5)

# شريط التقدم والحالة
progress = ttk.Progressbar(window, length=600, mode="determinate", style="TProgressbar")
progress.pack(pady=15)
status_label = tk.Label(window, text="الحالة / Status: جاهز / Ready", font=("Arial", 12), bg="#2E3B4E", fg="white")
status_label.pack(pady=5)


def update_progress_color(value):
    if value < 20:
        style.configure("TProgressbar", background="gray")
    elif value < 40:
        style.configure("TProgressbar", background="yellow")
    elif value < 60:
        style.configure("TProgressbar", background="orange")
    elif value < 80:
        style.configure("TProgressbar", background="red")
    elif value < 100:
        style.configure("TProgressbar", background="green")
    else:
        style.configure("TProgressbar", background="blue")


# --------------------------- دوال توليد المحتوى ---------------------------
def generate_text(prompt):
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"Error during generation: {e}")
        return None


def google_search(query, num_results=3):
    url = "https://www.googleapis.com/customsearch/v1"
    params = {"key": GOOGLE_SEARCH_API_KEY, "cx": GOOGLE_SEARCH_ENGINE_ID, "q": query, "num": num_results}
    try:
        response = requests.get(url, params=params)
        results = []
        if response.status_code == 200:
            data = response.json()
            items = data.get("items", [])
            for item in items:
                results.append({
                    "title": item.get("title"),
                    "link": item.get("link"),
                    "snippet": item.get("snippet")
                })
        else:
            messagebox.showerror("خطأ / Error", f"Google search failed: {response.status_code}")
        return results
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"Google search error: {e}")
        return []


def save_document(doc, title):
    try:
        filename = os.path.join(save_path, f"{title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx")
        doc.save(filename)
        return filename
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"Failed to save file: {e}")
        return None


def generate_article(topic, language):
    references = google_search(topic, num_results=3)
    if references:
        references_text = "\n\n".join(
            [f"Title: {ref['title']}\nLink: {ref['link']}\nSnippet: {ref['snippet']}" for ref in references])
    else:
        references_text = "لا توجد مراجع متاحة." if language == "العربية" else "No references available."
    if language == "العربية":
        prompt = (
            f"اكتب مقالاً احترافيًا ومميزًا عن '{topic}'. "
            "يجب أن يكون المقال شاملاً، يحوي تحليلاً عميقًا ورؤى فريدة، مع استخدام أسلوب أدبي راقٍ. "
            "اعتمد على المعلومات التالية المستخرجة من بحث جوجل كمرجع لتعزيز جودة المقال:\n\n"
            f"{references_text}\n\n"
            "رتب الأفكار بشكل منطقي وسلس وتأكد من إبراز النقاط الرئيسية بوضوح."
        )
    else:
        prompt = (
            f"Write a professional and exceptional article about '{topic}'. "
            "The article should be comprehensive, providing in-depth analysis and unique insights with an elegant writing style. "
            "Use the following Google search results as references to enhance the quality of the article:\n\n"
            f"{references_text}\n\n"
            "Ensure the ideas are well-organized and the key points are clearly highlighted."
        )
    return generate_text(prompt)


# --------------------------- العملية الرئيسية ---------------------------
def generate_and_save():
    topic = topic_entry.get().strip()
    language = language_var.get()
    selected_output = output_var.get()
    if not topic:
        messagebox.showerror("خطأ / Error", "يرجى إدخال عنوان الموضوع / Please enter the topic title.")
        return
    progress["value"] = 10
    update_progress_color(10)
    status_label.config(text="جارٍ توليد المقالة... / Generating article...")
    window.update_idletasks()
    article_text = generate_article(topic, language)
    if not article_text:
        status_label.config(text="فشل توليد المقالة / Article generation failed.")
        return
    # إزالة "##" من بداية السطر الأول إن وُجد:
    lines = article_text.splitlines()
    if lines and lines[0].startswith("##"):
        lines[0] = lines[0].lstrip("# ").strip()
        article_text = "\n".join(lines)
    progress["value"] = 50
    update_progress_color(50)
    window.update_idletasks()

    # إنشاء ملف وورد وتنسيق المحتوى وفق التعليمات الثابتة:
    doc = Document()
    # تحديد الإعدادات الثابتة بناءً على اللغة:
    if language == "العربية":
        main_title_font = "Amiri"
        body_font = "Amiri"
        main_title_alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        main_title_font = "Calibri"
        body_font = "Calibri"
        main_title_alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_alignment = WD_ALIGN_PARAGRAPH.LEFT

    # إضافة العنوان الرئيسي:
    main_title = doc.add_heading(topic, level=1)
    main_title.alignment = main_title_alignment
    for run in main_title.runs:
        run.font.name = main_title_font
        run.font.size = Pt(24)  # حجم العنوان الرئيسي
        run.font.color.rgb = RGBColor(0, 0, 255)  # العنوان باللون الأزرق

    # تقسيم المقال إلى أسطر؛ تحويل الأسطر التي تبدأ بـ "##" لعناوين فرعية:
    for line in article_text.splitlines():
        if line.strip().startswith("##"):
            subheading_text = line.strip().lstrip("#").strip()
            subheading = doc.add_heading(subheading_text, level=2)
            subheading.alignment = body_alignment
            for run in subheading.runs:
                run.font.name = body_font
                run.font.size = Pt(18)  # حجم العنوان الفرعي
                run.font.color.rgb = RGBColor(0, 0, 0)  # باللون الأسود
        else:
            para = doc.add_paragraph(line)
            para.alignment = body_alignment
            for run in para.runs:
                run.font.name = body_font
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

    insert_images(doc, topic)
    docx_filename = save_document(doc, topic)
    if not docx_filename:
        status_label.config(text="فشل حفظ ملف الوورد / Failed to save DOCX.")
        return
    if selected_output == "حفظ كملف وورد (DOCX)":
        output_file = docx_filename
    elif selected_output == "تحويل إلى PDF":
        output_file = convert_to_pdf(docx_filename)
        if output_file and doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    elif selected_output == "تحويل إلى صوت (MP3)":
        output_file = generate_audio(article_text, topic, language)
        if doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    elif selected_output == "تحويل إلى صوت (ElevenLabs)":
        api_key = elevenlabs_api_entry.get().strip()
        voice_id = elevenlabs_voice_entry.get().strip()
        if not api_key or not voice_id:
            messagebox.showerror("خطأ / Error", "يرجى إدخال ElevenLabs API Key و Voice ID.")
            return
        output_file = generate_audio_elevenlabs(article_text, topic, api_key, voice_id)
        if doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    progress["value"] = 100
    update_progress_color(100)
    window.update_idletasks()
    if output_file:
        status_label.config(text=f"تم الحفظ بنجاح / Successfully saved: {output_file}")
        messagebox.showinfo("نجاح / Success", f"تم إنشاء الملف:\n{output_file}")
    else:
        status_label.config(text="لم يتم إنشاء أي ملف.")


# زر بدء العملية
generate_button = tk.Button(window, text="توليد المقالة وحفظها / Generate and Save", command=generate_and_save,
                            font=("Arial", 14, "bold"), bg="#4CAF50", fg="white", padx=10, pady=5)
generate_button.pack(pady=20)

window.mainloop()

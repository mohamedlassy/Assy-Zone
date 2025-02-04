import google.generativeai as genai
from docx import Document
from tkinter import Tk, Label, Entry, Button, filedialog, messagebox, StringVar, Canvas
from tkinter.ttk import Progressbar, Style, Combobox
import os
from datetime import datetime

# تكوين مفتاح API
API_KEY = 'AIzaSyDvN7DZ6ETx-37bYnbWMNbFKWo2TO6lGjI'
genai.configure(api_key=API_KEY)

# إعداد نافذة البرنامج
window = Tk()
window.title("Script AI")
window.geometry("600x500")
window.resizable(False, False)
window.configure(bg="#2E3B4E")  # خلفية زرقاء داكنة أنيقة

# إعداد الأنماط لشريط التقدم
style = Style(window)
style.theme_use("default")
style.configure("TProgressbar", thickness=20)

# عنوان البرنامج
title_label = Label(window, text="Smart Writer", font=("Arial", 24, "bold"), fg="white", bg="#2E3B4E")
title_label.pack(pady=10)

# قائمة اختيار اللغة
language_label = Label(window, text="اختر اللغة / Choose Language:", font=("Arial", 14), bg="#2E3B4E", fg="white")
language_label.pack(pady=5)

language_var = StringVar()
language_var.set("العربية")  # اللغة الافتراضية
language_combobox = Combobox(window, textvariable=language_var, values=["العربية", "English"], state="readonly")
language_combobox.pack(pady=10)

# إدخال عنوان الموضوع
topic_label = Label(window, text="أدخل عنوان الموضوع / Enter the topic title:", font=("Arial", 14), bg="#2E3B4E", fg="white")
topic_label.pack(pady=5)
topic_entry = Entry(window, width=50)
topic_entry.pack(pady=10)

# المسار الافتراضي لحفظ الملفات
save_path = os.path.join(os.path.expanduser("~"), "Desktop")

# عرض المسار الحالي
save_path_label = Label(window, text=f"المسار الحالي / Current Path: {save_path}", font=("Arial", 10), bg="#2E3B4E", fg="white")
save_path_label.pack(pady=5)

# شريط التقدم
progress = Progressbar(window, length=400, mode="determinate", style="TProgressbar")
progress.pack(pady=10)

# تحديث لون شريط التقدم
def update_progress_color(value):
    if value < 20:
        style.configure("TProgressbar", troughcolor="white", background="gray")
    elif value < 40:
        style.configure("TProgressbar", troughcolor="white", background="yellow")
    elif value < 60:
        style.configure("TProgressbar", troughcolor="white", background="orange")
    elif value < 80:
        style.configure("TProgressbar", troughcolor="white", background="red")
    elif value < 100:
        style.configure("TProgressbar", troughcolor="white", background="green")
    else:
        style.configure("TProgressbar", troughcolor="white", background="blue")

# توليد النص باستخدام الذكاء الاصطناعي
def generate_text(prompt):
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"خطأ أثناء التوليد / Error during generation: {e}")
        return None

# إنشاء وحفظ المستند
def create_document(title, content):
    doc = Document()
    doc.add_heading(title.strip(), level=1)
    doc.add_paragraph(content.strip())
    return doc

def save_document(doc, title):
    try:
        filename = os.path.join(save_path, f"{title}_{datetime.now().strftime('%Y-%m-%d')}.docx")
        doc.save(filename)
        return filename
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"فشل حفظ الملف / Failed to save file: {e}")
        return None

# تنفيذ العملية الكاملة
def generate_and_save():
    topic = topic_entry.get()
    language = language_var.get()

    if not topic:
        messagebox.showerror("خطأ / Error", "يرجى إدخال عنوان الموضوع / Please enter the topic title.")
        return

    status_label.config(text="جارٍ التوليد... / Generating...")
    progress["value"] = 10
    update_progress_color(10)
    window.update_idletasks()

    if language == "العربية":
        main_story_prompt = f"اكتب موضوعًا شاملاً ومفصلًا عن {topic}. بدون مراجع أو إضافات أخرى."
    else:
        main_story_prompt = f"Write a comprehensive and detailed article about {topic}. No references or additional extras."

    content = generate_text(main_story_prompt)
    if not content:
        status_label.config(text="فشل التوليد / Generation failed.")
        return

    progress["value"] = 70
    update_progress_color(70)
    window.update_idletasks()

    doc = create_document(topic, content)

    progress["value"] = 90
    update_progress_color(90)
    window.update_idletasks()

    filename = save_document(doc, topic)
    if filename:
        progress["value"] = 100
        update_progress_color(100)
        status_label.config(text=f"تم الحفظ بنجاح في / Successfully saved in: {filename}")
        messagebox.showinfo("نجاح / Success", f"تم إنشاء الملف وحفظه بصيغة DOCX / File saved successfully as DOCX:\n{filename}")
    else:
        status_label.config(text="فشل الحفظ / Save failed.")

# زر توليد الموضوع
generate_button = Button(window, text="توليد الموضوع وحفظه / Generate and Save", command=generate_and_save, font=("Arial", 14), bg="#4CAF50", fg="white")
generate_button.pack(pady=20)

# حالة العملية
status_label = Label(window, text="الحالة / Status: جاهز / Ready", font=("Arial", 12), fg="white", bg="#2E3B4E")
status_label.pack(pady=10)

window.mainloop()

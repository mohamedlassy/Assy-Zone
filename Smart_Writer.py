import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import ttk, messagebox
import os
from datetime import datetime
import requests
from io import BytesIO

# مكتبات الميزات الجديدة:
from gtts import gTTS  # لتحويل النص إلى كلام باستخدام gTTS
from docx2pdf import convert  # لتحويل DOCX إلى PDF

# --------------------------- إعدادات وتكوينات API ---------------------------

# مفتاح Google Generative AI
GENAI_API_KEY = 'AIzaSyDvN7DZ6ETx-37bYnbWMNbFKWo2TO6lGjI'
genai.configure(api_key=GENAI_API_KEY)

# مفاتيح البحث:
GOOGLE_SEARCH_API_KEY = 'AIzaSyAN9T3joIAva5jWbUxAdvtgvhponcvEED4'
GOOGLE_SEARCH_ENGINE_ID = 'a048839a17d7b404c'  # معرّف محرك البحث المخصص

# مفتاح Pexels API (المستخدم لجلب الصور)
PEXELS_API_KEY = 'Afw2O4LTmWoNOps7spA9YhYAFbFVol8KfWjfSRxU059Wd8UqUXHRSAAJ'


# --------------------------- دوال البحث عن الصور ---------------------------

def get_pexels_images(query, num_images=2):
    """
    جلب صور من Pexels API حسب الموضوع المطلوب.
    """
    images = []
    url = "https://api.pexels.com/v1/search"
    headers = {"Authorization": PEXELS_API_KEY}
    params = {"query": query, "per_page": num_images}
    try:
        response = requests.get(url, headers=headers, params=params)
        if response.status_code == 200:
            data = response.json()
            photos = data.get("photos", [])
            for photo in photos:
                image_url = photo.get("src", {}).get("medium")
                if image_url:
                    img_response = requests.get(image_url)
                    if img_response.status_code == 200:
                        images.append(BytesIO(img_response.content))
        else:
            messagebox.showwarning("Pexels API Warning", f"Pexels API returned status code {response.status_code}")
    except Exception as e:
        messagebox.showerror("Pexels API Error", f"Error fetching images from Pexels: {e}")
    return images


def insert_images(doc, query):
    """
    يجلب الصور من Pexels ثم يُدرجها في المستند.
    """
    images = get_pexels_images(query, num_images=2)
    if images:
        doc.add_heading("Images / الصور", level=2)
        for img in images:
            try:
                doc.add_picture(img, width=Inches(5))
            except Exception as e:
                messagebox.showwarning("Image Insert Warning", f"Could not insert an image: {e}")


# --------------------------- دوال الميزات الجديدة ---------------------------

def generate_audio(text, title, language):
    """
    تحويل النص إلى كلام باستخدام gTTS وحفظه كملف MP3.
    """
    tts_lang = 'ar' if language == "العربية" else 'en'
    try:
        tts = gTTS(text=text, lang=tts_lang)
        # تجاهل "##" من العنوان:
        clean_title = title.replace("##", "")
        audio_filename = os.path.join(save_path, f"{clean_title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.mp3")
        tts.save(audio_filename)
        return audio_filename
    except Exception as e:
        messagebox.showerror("Audio Generation Error", f"Error generating audio with gTTS: {e}")
        return None


def generate_audio_elevenlabs(text, title, api_key, voice_id):
    """
    تحويل النص إلى كلام باستخدام ElevenLabs API وحفظه كملف MP3.
    """
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
    headers = {
        "xi-api-key": api_key,
        "Content-Type": "application/json"
    }
    payload = {
        "text": text,
        "voice_settings": {
            "stability": 0.5,
            "similarity_boost": 0.5
        }
    }
    try:
        response = requests.post(url, json=payload, headers=headers)
        if response.status_code == 200:
            # تجاهل "##" من العنوان:
            clean_title = title.replace("##", "")
            audio_filename = os.path.join(save_path,
                                          f"{clean_title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}_elevenlabs.mp3")
            with open(audio_filename, "wb") as f:
                f.write(response.content)
            return audio_filename
        else:
            messagebox.showerror("ElevenLabs Error",
                                 f"Error generating audio from ElevenLabs: {response.status_code}\n{response.text}")
            return None
    except Exception as e:
        messagebox.showerror("ElevenLabs Error", f"Exception during ElevenLabs audio generation: {e}")
        return None


def convert_to_pdf(docx_filename):
    """
    تحويل ملف DOCX إلى PDF باستخدام مكتبة docx2pdf.
    """
    try:
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        convert(docx_filename, pdf_filename)
        return pdf_filename
    except Exception as e:
        messagebox.showerror("PDF Conversion Error", f"Error converting to PDF: {e}")
        return None


# --------------------------- إعداد واجهة المستخدم ---------------------------

window = tk.Tk()
window.title("Smart Writer - برنامج الكتابة الذكية")
window.geometry("750x850")
window.resizable(False, False)
window.configure(bg="#2E3B4E")

# شريط القوائم
menu_bar = tk.Menu(window)
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Exit", command=window.quit)
menu_bar.add_cascade(label="File", menu=file_menu)
help_menu = tk.Menu(menu_bar, tearoff=0)
help_menu.add_command(label="About",
                      command=lambda: messagebox.showinfo("About", "Smart Writer\nDeveloped by Professional Developer"))
menu_bar.add_cascade(label="Help", menu=help_menu)
window.config(menu=menu_bar)

# إعداد نمط شريط التقدم
style = ttk.Style(window)
style.theme_use("clam")
style.configure("TProgressbar", thickness=20, background="#4CAF50", troughcolor="#D3D3D3")

# إطار الرأس (Header)
header_frame = tk.Frame(window, bg="#2E3B4E")
header_frame.pack(pady=20)
title_label = tk.Label(header_frame, text="Smart Writer", font=("Arial", 32, "bold"), fg="white", bg="#2E3B4E")
title_label.pack()
subtitle_label = tk.Label(header_frame, text="برنامج الكتابة الذكية", font=("Arial", 16), fg="lightgray", bg="#2E3B4E")
subtitle_label.pack()

# إطار الإدخالات (باستخدام pack لجميع الإطارات الفرعية)
input_frame = tk.Frame(window, bg="#2E3B4E")
input_frame.pack(pady=10, padx=20, fill="x")

# إطار اختيار اللغة
lang_frame = tk.Frame(input_frame, bg="#2E3B4E")
lang_frame.pack(anchor="w", pady=10)
lang_label = tk.Label(lang_frame, text="اختر اللغة / Choose Language:", font=("Arial", 14), bg="#2E3B4E", fg="white")
lang_label.pack(side="left", padx=(0, 10))
language_var = tk.StringVar(value="العربية")
lang_combobox = ttk.Combobox(lang_frame, textvariable=language_var, values=["العربية", "English"], state="readonly",
                             font=("Arial", 12), width=12)
lang_combobox.pack(side="left")

# إطار إدخال عنوان الموضوع
topic_frame = tk.Frame(input_frame, bg="#2E3B4E")
topic_frame.pack(anchor="w", pady=10)
topic_label = tk.Label(topic_frame, text="أدخل عنوان الموضوع / Enter Topic Title:", font=("Arial", 14), bg="#2E3B4E",
                       fg="white")
topic_label.pack(side="left", padx=(0, 10))
topic_entry = tk.Entry(topic_frame, width=40, font=("Arial", 12))
topic_entry.pack(side="left")

# إضافة قائمة سياقية (Context Menu) لحقل إدخال الموضوع للسماح باللصق عند النقر بزر الماوس الأيمن
def show_topic_context_menu(event):
    try:
        topic_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        topic_context_menu.grab_release()

topic_context_menu = tk.Menu(window, tearoff=0)
topic_context_menu.add_command(label="Paste", command=lambda: topic_entry.insert(tk.INSERT, window.clipboard_get()))
topic_entry.bind("<Button-3>", show_topic_context_menu)

# إطار اختيار نوع المخرجات عبر قائمة منسدلة
output_frame = tk.Frame(input_frame, bg="#2E3B4E")
output_frame.pack(anchor="w", pady=10)
output_label = tk.Label(output_frame, text="اختر نوع المخرجات / Select Output Type:", font=("Arial", 14), bg="#2E3B4E",
                        fg="white")
output_label.pack(side="left", padx=(0, 10))
# إضافة خيار ElevenLabs إلى القائمة
output_options = ["حفظ كملف وورد (DOCX)", "تحويل إلى PDF", "تحويل إلى صوت (MP3)", "تحويل إلى صوت (ElevenLabs)"]
output_var = tk.StringVar(value=output_options[0])
output_combobox = ttk.Combobox(output_frame, textvariable=output_var, values=output_options, state="readonly",
                               font=("Arial", 12), width=30)
output_combobox.pack(side="left")

# إطار بيانات ElevenLabs (يظهر فقط إذا اختار المستخدم خيار ElevenLabs)
elevenlabs_frame = tk.Frame(input_frame, bg="#2E3B4E")
# حقل إدخال API Key
elevenlabs_api_label = tk.Label(elevenlabs_frame, text="ElevenLabs API Key:", font=("Arial", 12), bg="#2E3B4E",
                                fg="white")
elevenlabs_api_label.pack(side="left", padx=(0, 5))
elevenlabs_api_entry = tk.Entry(elevenlabs_frame, width=30, font=("Arial", 12))
elevenlabs_api_entry.pack(side="left", padx=(0, 10))

# إضافة قائمة سياقية لحقل API Key
def show_api_context_menu(event):
    try:
        api_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        api_context_menu.grab_release()

api_context_menu = tk.Menu(window, tearoff=0)
api_context_menu.add_command(label="Paste",
                             command=lambda: elevenlabs_api_entry.insert(tk.INSERT, window.clipboard_get()))
elevenlabs_api_entry.bind("<Button-3>", show_api_context_menu)

# حقل إدخال Voice ID
elevenlabs_voice_label = tk.Label(elevenlabs_frame, text="Voice ID:", font=("Arial", 12), bg="#2E3B4E", fg="white")
elevenlabs_voice_label.pack(side="left", padx=(0, 5))
elevenlabs_voice_entry = tk.Entry(elevenlabs_frame, width=20, font=("Arial", 12))
elevenlabs_voice_entry.pack(side="left")

# إضافة قائمة سياقية لحقل Voice ID
def show_voice_context_menu(event):
    try:
        voice_context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        voice_context_menu.grab_release()

voice_context_menu = tk.Menu(window, tearoff=0)
voice_context_menu.add_command(label="Paste",
                               command=lambda: elevenlabs_voice_entry.insert(tk.INSERT, window.clipboard_get()))
elevenlabs_voice_entry.bind("<Button-3>", show_voice_context_menu)

# إطار قائمة الخيارات الخاصة بملف الوورد الوسيط (لإبقاء الملف أو حذفه)
doc_option_frame = tk.Frame(input_frame, bg="#2E3B4E")
doc_option_label = tk.Label(doc_option_frame, text="ملف الوورد الوسيط:", font=("Arial", 14), bg="#2E3B4E", fg="white")
doc_option_label.pack(side="left", padx=(0, 10))
doc_option_var = tk.StringVar(value="حذف ملف الوورد")
doc_option_combobox = ttk.Combobox(doc_option_frame, textvariable=doc_option_var,
                                   values=["حذف ملف الوورد", "الابقاء على ملف الوورد"], state="readonly",
                                   font=("Arial", 12), width=30)
doc_option_combobox.pack(side="left")
# لن يتم عرضها تلقائيًا؛ ستظهر بناءً على اختيار المخرجات (PDF أو صوت)

# تحديث ظهور حقول ElevenLabs وقائمة خيار ملف الوورد بناءً على اختيار نوع المخرجات
def update_output_options(*args):
    # تحديث ظهور حقل ElevenLabs:
    if output_var.get() == "تحويل إلى صوت (ElevenLabs)":
        elevenlabs_frame.pack(anchor="w", pady=5)
    else:
        elevenlabs_frame.pack_forget()
    # إظهار أو إخفاء قائمة خيار ملف الوورد الوسيط:
    if output_var.get() in ["تحويل إلى PDF", "تحويل إلى صوت (MP3)", "تحويل إلى صوت (ElevenLabs)"]:
        doc_option_frame.pack(anchor="w", pady=5)
    else:
        doc_option_frame.pack_forget()

output_var.trace_add("write", update_output_options)

# تحديد مسار الحفظ ليكون سطح المكتب
save_path = os.path.join(os.path.expanduser("~"), "Desktop")

save_path_info = tk.Label(window, text=f"سيتم الحفظ في: {save_path}", font=("Arial", 10), bg="#2E3B4E", fg="white")
save_path_info.pack(pady=5)

# شريط التقدم
progress = ttk.Progressbar(window, length=600, mode="determinate", style="TProgressbar")
progress.pack(pady=15)
# تسمية الحالة
status_label = tk.Label(window, text="الحالة / Status: جاهز / Ready", font=("Arial", 12), bg="#2E3B4E", fg="white")
status_label.pack(pady=5)

def update_progress_color(value):
    """تحديث لون شريط التقدم تبعًا للقيمة."""
    if value < 20:
        style.configure("TProgressbar", background="gray")
    elif value < 40:
        style.configure("TProgressbar", background="yellow")
    elif value < 60:
        style.configure("TProgressbar", background="orange")
    elif value < 80:
        style.configure("TProgressbar", background="red")
    elif value < 100:
        style.configure("TProgressbar", background="green")
    else:
        style.configure("TProgressbar", background="blue")


# --------------------------- دوال توليد المحتوى ---------------------------

def generate_text(prompt):
    """
    دالة لتوليد نص باستخدام الذكاء الاصطناعي عبر Google Generative AI.
    """
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"Error during generation: {e}")
        return None


def google_search(query, num_results=3):
    """
    دالة للبحث في جوجل باستخدام Google Custom Search API.
    """
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": GOOGLE_SEARCH_API_KEY,
        "cx": GOOGLE_SEARCH_ENGINE_ID,
        "q": query,
        "num": num_results
    }
    try:
        response = requests.get(url, params=params)
        results = []
        if response.status_code == 200:
            data = response.json()
            items = data.get("items", [])
            for item in items:
                results.append({
                    "title": item.get("title"),
                    "link": item.get("link"),
                    "snippet": item.get("snippet")
                })
        else:
            messagebox.showerror("خطأ / Error", f"Google search failed: {response.status_code}")
        return results
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"Google search error: {e}")
        return []


def save_document(doc, title):
    """
    دالة لحفظ المستند بصيغة DOCX في مسار الحفظ.
    """
    try:
        filename = os.path.join(save_path, f"{title}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx")
        doc.save(filename)
        return filename
    except Exception as e:
        messagebox.showerror("خطأ / Error", f"Failed to save file: {e}")
        return None


def generate_article(topic, language):
    """
    توليد مقال مفصل عن الموضوع باستخدام نتائج بحث جوجل كمراجع، بصياغة احترافية وفريدة.
    """
    references = google_search(topic, num_results=3)
    if references:
        references_text = "\n\n".join(
            [f"Title: {ref['title']}\nLink: {ref['link']}\nSnippet: {ref['snippet']}" for ref in references]
        )
    else:
        references_text = "لا توجد مراجع متاحة." if language == "العربية" else "No references available."

    if language == "العربية":
        prompt = (
            f"اكتب مقالاً احترافيًا ومميزًا عن '{topic}'. "
            "يجب أن يكون المقال شاملاً، يحوي تحليلاً عميقًا ورؤى فريدة، مع استخدام أسلوب أدبي راقٍ. "
            "اعتمد على المعلومات التالية المستخرجة من بحث جوجل كمرجع لتعزيز جودة المقال:\n\n"
            f"{references_text}\n\n"
            "رتب الأفكار بشكل منطقي وسلس وتأكد من إبراز النقاط الرئيسية بوضوح."
        )
    else:
        prompt = (
            f"Write a professional and exceptional article about '{topic}'. "
            "The article should be comprehensive, providing in-depth analysis and unique insights with an elegant writing style. "
            "Use the following Google search results as references to enhance the quality of the article:\n\n"
            f"{references_text}\n\n"
            "Ensure the ideas are well-organized and the key points are clearly highlighted."
        )
    return generate_text(prompt)


# --------------------------- العملية الرئيسية ---------------------------

def generate_and_save():
    topic = topic_entry.get().strip()
    language = language_var.get()
    selected_output = output_var.get()  # نوع المخرجات المطلوب من القائمة

    if not topic:
        messagebox.showerror("خطأ / Error", "يرجى إدخال عنوان الموضوع / Please enter the topic title.")
        return

    progress["value"] = 10
    update_progress_color(10)
    status_label.config(text="جارٍ توليد المقالة... / Generating article...")
    window.update_idletasks()

    article_text = generate_article(topic, language)
    if not article_text:
        status_label.config(text="فشل توليد المقالة / Article generation failed.")
        return

    # إزالة "##" من بداية السطر الأول إن وُجد:
    lines = article_text.splitlines()
    if lines and lines[0].startswith("##"):
        lines[0] = lines[0].lstrip("# ").strip()
        article_text = "\n".join(lines)

    progress["value"] = 50
    update_progress_color(50)
    window.update_idletasks()

    output_file = None
    docx_filename = None  # الملف الوسيط

    # إنشاء ملف وورد كخطوة وسيطة دائمًا
    doc = Document()
    doc.add_heading(topic, level=1)
    doc.add_paragraph(article_text)
    insert_images(doc, topic)
    docx_filename = save_document(doc, topic)
    if not docx_filename:
        status_label.config(text="فشل حفظ ملف الوورد / Failed to save DOCX.")
        return

    # اعتمادًا على نوع المخرجات المطلوب:
    if selected_output == "حفظ كملف وورد (DOCX)":
        output_file = docx_filename
        # لا يتم حذف الملف الوسيط
    elif selected_output == "تحويل إلى PDF":
        output_file = convert_to_pdf(docx_filename)
        if output_file and doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    elif selected_output == "تحويل إلى صوت (MP3)":
        output_file = generate_audio(article_text, topic, language)
        if doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")
    elif selected_output == "تحويل إلى صوت (ElevenLabs)":
        # الحصول على بيانات ElevenLabs من الحقول
        api_key = elevenlabs_api_entry.get().strip()
        voice_id = elevenlabs_voice_entry.get().strip()
        if not api_key or not voice_id:
            messagebox.showerror("خطأ / Error", "يرجى إدخال ElevenLabs API Key و Voice ID.")
            return
        output_file = generate_audio_elevenlabs(article_text, topic, api_key, voice_id)
        if doc_option_var.get() == "حذف ملف الوورد":
            try:
                os.remove(docx_filename)
            except Exception as e:
                messagebox.showwarning("Warning", f"Unable to remove intermediate DOCX file: {e}")

    progress["value"] = 100
    update_progress_color(100)
    window.update_idletasks()

    if output_file:
        status_label.config(text=f"تم الحفظ بنجاح / Successfully saved: {output_file}")
        messagebox.showinfo("نجاح / Success", f"تم إنشاء الملف:\n{output_file}")
    else:
        status_label.config(text="لم يتم إنشاء أي ملف.")


# زر بدء العملية
generate_button = tk.Button(window, text="توليد المقالة وحفظها / Generate and Save", command=generate_and_save,
                            font=("Arial", 14), bg="#4CAF50", fg="white", padx=10, pady=5)
generate_button.pack(pady=20)

# بدء الحلقة الرئيسية للتطبيق
window.mainloop()
